"""졸업논문 .docx 조립 — paper/*.md + 그림 + 자동 생성 표 → 졸업논문_초안v1.docx

학교 양식(붙임1_졸업논문양식.pdf)의 순서를 그대로 따른다:
  제목(국문·영문) → Abstract(국문요약 500자 이내) → Keyword(5개 이하)
  → Ⅰ. Introduction → Ⅱ. Related Works → Ⅲ. Proposed Method
  → Ⅳ. Experimental Results → Ⅴ. Conclusions → 참고문헌

**원고는 마크다운이 원본이다.** 이 스크립트를 다시 돌리면 .docx가 새로 만들어지므로
.docx를 손으로 고치면 다음 실행에서 사라진다. 고칠 것은 paper/*.md 쪽이다.

실행: python -m src.report.make_thesis_docx
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.report import docx_build as B
from src.report import md_to_docx as M
from src.report.paper_tables import all_tables

ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
OUT = ROOT / "졸업논문_초안v1.docx"

TITLE_KO = "행동에 비용이 붙을 때 학습은 언제 규칙을 이기는가: 행동 비용 λ에 따른 성능 지도와 임계 비용"
TITLE_EN = "When does learning beat a fixed rule under action cost? A performance map over the action cost λ"
AUTHOR_LINE = "○○학과 20○○○○○○ 전혜성   [본인 확인 필요 — 학과·학번을 채울 것]"

CHAPTERS = [
    ("paper/01_서론.md", {}),
    ("paper/02_관련연구.md", {}),
    ("paper/03_방법.md", {}),
    ("paper/04_결과.md", {}),
    ("paper/05_결론.md", {}),
]

# 서지는 thesis-writing 스킬의 확인된 목록을 그대로 옮긴 것이다.
# **확인되지 않은 항목(권·페이지 등)은 넣지 않는다.** 지어낸 서지는 논문에서 가장 나쁜 종류의 오류다.
#
# 2026-08-29 확인: 비어 있던 [2]·[3]의 페이지 번호를 DBLP에서 확인해 채웠다.
#   [2] https://dblp.org/pid/176/5465.html → "AAMAS 2022: 669-677"
#       IFAAMAS 원문 파일명 p669.pdf 로 시작 쪽이 교차 확인됨
#       (https://ifaamas.org/Proceedings/aamas2022/pdfs/p669.pdf)
#   [3] https://dblp.org/rec/conf/nips/AgarwalSCCB21.html → "NeurIPS 2021: 29304-29320"
# 제출 전 본인이 원문으로 한 번 더 대조할 것 — 여기 적힌 것은 서지 데이터베이스 기준이다.
REFERENCES = [
    'A. Biedenkapp, R. Rajan, F. Hutter, and M. Lindauer, "TempoRL: Learning when to act," '
    'Proceedings of the 38th International Conference on Machine Learning (ICML), PMLR Vol.139, '
    'pp.914-924, 2021.',
    'A. Jacq, J. Ferret, O. Pietquin, and M. Geist, "Lazy-MDPs: Towards interpretable '
    'reinforcement learning by learning when to act," Proceedings of the 21st International '
    'Conference on Autonomous Agents and Multiagent Systems (AAMAS), pp.669-677, 2022.',
    'R. Agarwal, M. Schwarzer, P. S. Castro, A. Courville, and M. G. Bellemare, "Deep '
    'reinforcement learning at the edge of the statistical precipice," Advances in Neural '
    'Information Processing Systems (NeurIPS), Vol.34, pp.29304-29320, 2021.',
]

# 본문 폭은 16cm(A4 21 − 여백 2.5×2)다. 그림을 그보다 좁게 넣으면 그림 안의 글자도
# 같은 비율로 작아진다 — 14cm면 범례가 6pt가 되어 본문 10pt보다 작아진다.
# 논문의 핵심 근거인 λ 지도가 가장 읽기 어려워서는 안 되므로 전부 본문 폭에 맞춘다.
FIG_WIDTH = {"fig1": 16.0, "fig2": 16.0, "fig3": 16.0, "fig4": 16.0, "fig5": 16.0,
             "fig6": 16.0}


def _read(path: str) -> str | None:
    p = ROOT / path
    if not p.exists():
        print(f"  [빠짐] {path} 없음 — 이 장은 건너뛴다")
        return None
    return p.read_text(encoding="utf-8")


def title_block(doc) -> None:
    B.para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, after=4, first_indent=False)
    p = B.para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.25, after=4, first_indent=False)
    B.set_run_font(p.add_run(TITLE_KO), B.HEAD_FONT, 15, bold=True)
    p = B.para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.25, after=8, first_indent=False)
    B.set_run_font(p.add_run(TITLE_EN), B.HEAD_FONT, 11.5, italic=True)
    p = B.para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.2, after=12, first_indent=False)
    B.set_run_font(p.add_run(AUTHOR_LINE), B.HEAD_FONT, 10, color="B3261E")


def abstract_block(doc) -> None:
    md = _read("paper/00_초록.md")
    if md is None:
        return
    body, keywords = "", ""
    for line in md.splitlines():
        s = line.strip()
        if not s or s.startswith("#") or s.startswith("<!--") or s.startswith(">"):
            continue
        if s.startswith("Keyword"):
            keywords = s
        else:
            body += (" " if body else "") + s
    p = B.para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.3, before=6, after=4,
               first_indent=False)
    B.set_run_font(p.add_run("Abstract"), B.HEAD_FONT, 12, bold=True)
    n = len(re.sub(r"\s", "", body))
    if n > 500:
        print(f"  ⚠ 국문요약이 {n}자 — 양식은 500자 이내다. 줄일 것")
    else:
        print(f"  국문요약 {n}자 (500자 이내 — 통과)")
    B.para(doc, body, line=1.7, after=6)
    if keywords:
        p = B.para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.3, after=10, first_indent=False)
        B.set_run_font(p.add_run(keywords), B.BODY_FONT, 10, bold=True)


def references_block(doc) -> None:
    B.add_chapter_heading(doc, "참 고 문 헌 (References)")
    for i, r in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.left_indent = B.Cm(0.9)
        pf.first_line_indent = B.Cm(-0.9)
        pf.space_after = B.Pt(3)
        B.set_run_font(p.add_run(f"[{i}] "), B.BODY_FONT, 9.5)
        B.set_run_font(p.add_run(r), B.BODY_FONT, 9.5)


def build() -> Path:
    print("졸업논문 .docx 조립 시작")
    tables = all_tables(fairness_variants=_fairness_variants())
    print(f"  자동 생성 표: {', '.join(tables)}")
    doc = B.new_document()
    title_block(doc)
    abstract_block(doc)
    counter = M.new_counter()
    for path, _ in CHAPTERS:
        md = _read(path)
        if md is None:
            continue
        M.render(doc, md, auto_tables=tables, fig_width=FIG_WIDTH, counter=counter)
    references_block(doc)
    doc.save(OUT)
    n_par = len(doc.paragraphs)
    n_tab = len(doc.tables)
    n_img = len(doc.inline_shapes)
    print(f"  저장: {OUT.name} | 문단 {n_par} · 표 {n_tab}(번호 {counter['tab']}) · "
          f"그림 {n_img}(번호 {counter['fig']})")
    return OUT


def _fairness_variants() -> dict[str, str]:
    """변종 방이 실제로 만들어져 있으면 공정성 표(표 3)에 자동으로 줄이 늘어난다."""
    labels = {
        "MountainCar-v0@budget1M_epsconst": "예산 1M (ε=0.2 고정)",
        "MountainCar-v0@budget1M_epsdecay": "예산 1M + ε 감소",
        "MountainCar-v0@budget1M_wide": "예산 1M + ε 감소 + 용량 확대",
    }
    return {k: v for k, v in labels.items() if (ROOT / "results" / "aggregate" / f"{k}_iqm.csv").exists()}


if __name__ == "__main__":
    build()
