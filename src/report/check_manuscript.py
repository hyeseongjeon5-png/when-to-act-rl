"""원고 일관성 점검 — 실험이 갱신되면서 낡아버린 문장을 찾는다.

논문의 가장 흔한 사고는 '숫자를 고쳤는데 그 숫자를 언급한 다른 문장을 안 고친 것'이다.
자동 생성되는 Ⅳ장과 달리 사람이 쓴 장(Ⅰ·Ⅱ·Ⅲ·Ⅴ·초록)은 손으로 고쳐야 하므로
여기서 위험한 표현을 기계적으로 훑는다. **자동으로 고치지 않는다** — 사람이 보고 판단할 목록만 만든다.

실행: python -m src.report.check_manuscript
"""
from __future__ import annotations

import json
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
AGG = ROOT / "results" / "aggregate"

# (이름, 정규식, 왜 위험한가)
RISKS = [
    ("환경 수 표현", r"두 환경|환경이 둘|양쪽 환경",
     "환경이 셋이 되면 낡는다 — 개수를 세는 표현은 확인할 것"),
    ("낡은 pump 값", r"120\.5|119\.8",
     "pump 규칙의 현재 값은 results/aggregate에서 확인할 것"),
    ("λ* 하드코딩", r"λ\?\*\s*(=|는)\s*\d",
     "λ*는 실험이 갱신되면 바뀐다 — 본문 대신 표를 가리킬 것"),
    ("기여 개수", r"얻은 것은 (두|세|네|다섯) 가지|기여는 (두|세|네) 가지",
     "항목을 늘리면 개수 표현이 낡는다"),
    ("확정 어투", r"확인했다|입증했다|증명했다",
     "실제로 확인한 범위를 넘지 않는지 볼 것"),
    ("결과 대기 표시", r"\[실험 대기 — 결과 나오면 확정\]",
     "그 실험이 끝나면 이 문장을 결과로 채우고 표시를 지울 것"),
    ("본인 확인 표시", r"\[본인 확인 필요[^\]]*\]",
     "제출 전 본인이 채워야 하는 자리"),
]


def numbers_in_use() -> dict:
    """집계 파일의 현재 기준 수치 — 원고와 대조할 때 쓴다."""
    out = {}
    for p in sorted(AGG.glob("*_lambda_star.json")):
        try:
            d = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
        env = d.get("env_id", p.stem)
        out[env] = {r["learner"]: {"CI": r.get("lam_star_ci"), "점추정": r.get("lam_star_pt")}
                    for r in d.get("results_vs_best_rule", d.get("results", []))}
    return out


def section_refs() -> list[str]:
    """원고가 가리키는 절 번호가 실제로 있는지 본다.

    왜 필요한가: 절을 하나 끼워 넣으면 뒤 번호가 밀리는데, **그 절을 가리키던 다른 장의 문장은
    그대로 남는다.** 실제로 Ⅲ장에 기준선 감사를 6.2로 넣자 인과 검사가 6.3이 되었고,
    Ⅴ장이 여전히 '6.2절'을 가리키고 있었다. 읽는 사람은 엉뚱한 절로 간다.
    """
    heads = {}
    for f in sorted(PAPER.glob("0[0-9]_*.md")):
        ch = None
        for line in f.read_text(encoding="utf-8").splitlines():
            m = re.match(r"^# ([ⅠⅡⅢⅣⅤ])\.", line.strip())
            if m:
                ch = m.group(1)
            m2 = re.match(r"^#{2,3}\s*(\d+(?:\.\d+)?)[\.\s]", line.strip())
            if m2 and ch:
                heads.setdefault(ch, set()).add(m2.group(1))
    out, bad = [], 0
    for f in sorted(PAPER.glob("0[0-9]_*.md")):
        text = f.read_text(encoding="utf-8")
        for m in re.finditer(r"([ⅠⅡⅢⅣⅤ])장\s*(\d+(?:\.\d+)?)절", text):
            ch, sec = m.group(1), m.group(2)
            if sec not in heads.get(ch, set()):
                ln = text[:m.start()].count(chr(10)) + 1
                out.append(f"  [없는 절] {f.name} {ln}행: '{ch}장 {sec}절' — "
                           f"{ch}장에 있는 절: {sorted(heads.get(ch, []))}")
                bad += 1
    if not out:
        n = sum(len(v) for v in heads.values())
        out = [f"  [맞음] 원고가 가리키는 절 번호가 모두 실제로 있다 (절 {n}개 확인)"]
    return out


def numbering_rules() -> list[str]:
    """조립된 .docx에서 그림·표 번호가 지시받은 배치와 맞는지 본다.

    왜 필요한가: 번호는 등장 순서로 자동 부여되므로 **다른 장에 표를 하나 추가하면
    조용히 밀린다.** 실제로 Ⅲ장에 표를 넣었더니 '표 1 = 임계 비용 λ*'가 표 2로 밀렸다.
    사람 눈에는 잘 안 띄는 종류의 어긋남이라 기계가 본다.
    """
    docx = ROOT / "졸업논문_초안v1.docx"
    if not docx.exists():
        return ["  [건너뜀] .docx가 아직 없다"]
    try:
        import mammoth
        with docx.open("rb") as f:
            html = mammoth.convert_to_html(f).value
    except Exception as e:
        return [f"  [건너뜀] .docx를 읽지 못함: {e}"]
    text = re.sub(r"<[^>]+>", " ", html)
    caps = {}
    for m in re.finditer(r"(표|그림) (\d+)\.\s*([^|]{0,60})", text):
        key = f"{m.group(1)} {m.group(2)}"
        caps.setdefault(key, re.sub(r"\s+", " ", m.group(3)).strip())
    # 지시받은 배치 (사용자 요구사항)
    want = [("표 1", "임계 비용"), ("표 2", "실험 설정"),
            ("그림 1", "래퍼"), ("그림 2", "λ-성능 지도"), ("그림 3", "λ-성능 지도"),
            ("그림 4", "무행동 붕괴")]
    out = []
    for key, must in want:
        got = caps.get(key)
        if got is None:
            out.append(f"  [없음] {key} — 아직 만들어지지 않았다 (실험이 끝나면 생길 수 있다)")
        elif must not in got:
            out.append(f"  [어긋남] {key}가 '{got[:40]}' 이다 — '{must}'이어야 한다")
        else:
            out.append(f"  [맞음] {key}. {got[:46]}")
    return out


def check_dangling_headings() -> None:
    """제목 바로 뒤에 또 제목이 오는 곳을 찾는다.

    그런 절은 "무엇을 말하는 절인지" 한 문장도 없이 시작한다. 실제로 이 원고의
    Ⅳ장 2절(가장 중요한 절)이 아무 문장 없이 표부터 시작하고 있었다 —
    심사자가 표를 스스로 해석해야 했다. 이 저장소의 글쓰기 규칙은 "결론을 맨 위에 쓴다"이다.
    """
    print(chr(10) + "절이 문장 없이 시작하지 않는가")
    bad = []
    for f in sorted(PAPER.glob("*.md")):
        if f.name.startswith("00_양식"):
            continue
        lines = f.read_text(encoding="utf-8").split(chr(10))
        for i, l in enumerate(lines):
            if not re.match(r"^#{2,3} ", l):
                continue
            for j in range(i + 1, len(lines)):
                t = lines[j].strip()
                if not t:
                    continue
                if re.match(r"^#{2,3} ", t):
                    bad.append((f.name, i + 1, l.strip(), t))
                break
    if not bad:
        print("  [맞음] 모든 절이 문장이나 목록으로 시작한다")
        return
    for fn_, ln, h, nxt in bad:
        print(f"  [고칠 것] {fn_}:{ln} '{h[:38]}' 바로 뒤에 '{nxt[:30]}'")
        print("        └ 이 절이 무엇을 말하는지 한 문장을 앞에 둘 것")


# 고유명사는 대문자 규칙의 예외다 (환경 이름·기법 이름·학회명·장 제목)
PROPER_NOUNS = {
    "MountainCar", "LunarLander", "MinAtar", "Freeway", "DQN", "TempoRL", "Lazy", "MDP", "MDPs",
    "IQM", "CI", "NeurIPS", "ICML", "AAMAS", "PMLR", "Gymnasium", "Adam", "Double", "Fig", "Table",
    "Proceedings", "International", "Conference", "Machine", "Learning", "Advances", "Neural",
    "Information", "Processing", "Systems", "Autonomous", "Agents", "Multiagent",
    "Introduction", "Related", "Works", "Proposed", "Method", "Experimental", "Results",
    "Conclusions", "Abstract", "Keyword", "When", "Critical", "Summary",
}


def check_citation_order() -> None:
    """양식: 참고문헌은 **본문 인용 순서대로** 번호를 매긴다.

    번호를 손으로 붙이므로, 문단을 옮기거나 문헌을 추가하면 조용히 어긋난다.
    사람 눈으로는 잘 안 보이고 심사에서는 바로 지적당하는 종류다.
    """
    print(chr(10) + "참고문헌 번호가 인용 순서와 맞는가")
    order, seen = [], set()
    for name in ("01_서론", "02_관련연구", "03_방법", "04_결과", "05_결론"):
        f = PAPER / (name + ".md")
        if not f.exists():
            continue
        for m in re.finditer(r"\[(\d+)\]", f.read_text(encoding="utf-8")):
            n = int(m.group(1))
            if n not in seen:
                seen.add(n)
                order.append((n, name))
    nums = [n for n, _ in order]
    if nums == sorted(nums):
        print(f"  [맞음] 본문에 나오는 순서가 {nums} 로 번호와 일치한다")
    else:
        print(f"  [고칠 것] 본문에 나오는 순서가 {nums} 다 — 번호를 다시 매길 것")
        for n, where in order:
            print(f"        [{n}] 처음 등장: {where}")


def check_english_case() -> None:
    """양식: 영문은 문장의 첫 자만 대문자, 나머지는 소문자 (고유명사 제외).

    영어 제목을 Title Case로 쓰는 습관 때문에 자주 어긋난다.
    """
    print(chr(10) + "영문 제목이 '첫 자만 대문자' 규칙을 지키는가")
    bad = []
    targets = list(PAPER.glob("*.md")) + [ROOT / "src" / "report" / "paper_tables.py"]
    caps = ROOT / "results" / "figures" / "paper" / "captions.json"
    texts = []
    for f in targets:
        if f.name.startswith("00_양식") or not f.exists():
            continue
        for m in re.finditer(r"TABCAP:[^|]*\|\s*([^>\"']*?)\s*(?:-->|\")",
                             f.read_text(encoding="utf-8"), re.S):
            texts.append((f.name, " ".join(m.group(1).split())))
    if caps.exists():
        for k, v in json.loads(caps.read_text(encoding="utf-8")).items():
            texts.append((k, v.get("en", "")))
    for where, en in texts:
        if not en or not re.search(r"[A-Za-z]", en):
            continue
        words = re.findall("[A-Za-z][A-Za-z-]*", en)
        hits = [w for i, w in enumerate(words)
                if i > 0 and w[0].isupper() and w not in PROPER_NOUNS
                and not w.startswith(("Mount", "Lunar", "Min"))]
        if hits:
            bad.append((where, hits, en))
    if not bad:
        print(f"  [맞음] 영문 제목 {len(texts)}개 모두 규칙을 지킨다")
        return
    for where, hits, en in bad:
        print(f"  [고칠 것] {where}: {', '.join(hits)}")
        print(f"        └ {en[:64]}")


def check_source_control_chars() -> None:
    """소스에 제어문자가 섞여 들어가지 않았는지 본다.

    2026-08-29 사고: 정규식의 낱말 경계 \\b 가 편집 과정에서 **백스페이스 문자(0x08)** 로
    바뀌어 있었다. 눈으로는 똑같아 보이고 sed/cat 출력에도 안 보이는데, 정규식이
    아무것도 못 찾아 검사가 **항상 통과**했다. 헛도는 검사는 없는 검사보다 나쁘다 —
    "확인했다"는 착각을 주기 때문이다.
    """
    print(chr(10) + "검사 도구 자체가 성한가 (소스에 제어문자)")
    bad_codes = {7: "\\a 벨", 8: "\\b 백스페이스",
                 11: "\\v 세로탭", 12: "\\f 폼피드", 27: "escape"}
    hits = []
    roots = [ROOT / "src", ROOT / "paper", ROOT / "docs", ROOT / "experiments"]
    for root in roots:
        if not root.exists():
            continue
        for f in sorted(root.rglob("*")):
            if f.suffix not in (".py", ".md", ".yaml", ".yml", ".sh") or not f.is_file():
                continue
            try:
                text = f.read_text(encoding="utf-8")
            except Exception:
                continue
            for ln, line in enumerate(text.split(chr(10)), 1):
                for ch in line:
                    if ord(ch) in bad_codes:
                        hits.append((f.relative_to(ROOT), ln, bad_codes[ord(ch)]))
                        break
    if not hits:
        print("  [맞음] 검사 대상 소스에 제어문자가 없다")
        return
    for f, ln, what in hits:
        print(f"  [고칠 것] {f}:{ln} — {what} 문자가 들어 있다")
        print("        └ 정규식이라면 아무것도 못 찾고 조용히 통과했을 수 있다")


ORDINALS = ["첫", "둘", "셋", "넷", "다섯", "여섯", "일곱", "여덟", "아홉", "열"]


def check_ordinal_order() -> None:
    """'첫째·둘째·셋째…'가 순서대로 나오는지 본다.

    2026-08-29에 실제로 Ⅴ장이 첫째 → 둘째 → 셋째 → **다섯째 → 넷째** 순으로 나와 있었다.
    문단을 옮기다 생긴 것으로, 눈으로 훑으면 잘 안 보이는데 심사자는 바로 본다.
    """
    print(chr(10) + "'첫째·둘째…'가 순서대로 나오는가")
    bad = []
    for f in sorted(PAPER.glob("*.md")):
        if f.name.startswith("00_양식"):
            continue
        seen = []
        for ln, line in enumerate(f.read_text(encoding="utf-8").split(chr(10)), 1):
            t = line.strip().lstrip("*").strip()
            for i, o in enumerate(ORDINALS):
                if t.startswith(o + "째"):
                    seen.append((i + 1, ln))
                    break
        nums = [n for n, _ in seen]
        if nums and nums != sorted(nums):
            bad.append((f.name, seen))
    if not bad:
        print("  [맞음] 차례를 세는 표현이 모두 순서대로 나온다")
        return
    for name, seen in bad:
        order = " → ".join(ORDINALS[n - 1] + "째" for n, _ in seen)
        print(f"  [고칠 것] {name}: {order}")
        print(f"        └ 나온 줄: {[ln for _, ln in seen]}")


def check_caption_position() -> None:
    """양식: **그림 제목은 그림 하단, 표 제목은 표 상단.** 조립된 .docx에서 직접 확인한다.

    지금은 코드가 그렇게 만들도록 짜여 있지만, 코드가 그렇다는 것과 결과물이 그렇다는 것은
    다르다. 만들어진 파일을 열어 순서를 본다.
    """
    print(chr(10) + "그림 제목은 아래, 표 제목은 위에 있는가 (.docx 확인)")
    docx = ROOT / "졸업논문_초안v1.docx"
    if not docx.exists():
        print("  [건너뜀] .docx가 없다 — 먼저 make_thesis_docx를 돌릴 것")
        return
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except Exception as e:
        print(f"  [확인 못 함] python-docx를 못 불러왔다: {e}")
        return
    doc = Document(str(docx))
    body = doc.element.body
    # 본문 순서대로 (종류, 내용) 목록을 만든다. 표는 하나의 항목으로 센다.
    seq = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            txt = "".join(n.text or "" for n in child.iter(qn("w:t"))).strip()
            kind = "그림파일" if child.find(".//" + qn("a:blip")) is not None else "글"
            seq.append((kind, txt))
        elif child.tag == qn("w:tbl"):
            seq.append(("표", ""))
    bad = 0
    for i, (kind, txt) in enumerate(seq):
        if kind == "그림파일":
            near = [t for k, t in seq[i + 1:i + 4] if k == "글" and t]
            if not any(re.match(r"그림 \d+\.", t) for t in near):
                print(f"  [고칠 것] {i}번째 그림 아래에 '그림 n.' 제목이 없다 (뒤에 온 글: {near[:2]})")
                bad += 1
        elif kind == "표":
            near = [t for k, t in seq[max(0, i - 3):i] if k == "글" and t]
            if not any(re.match(r"표 \d+\.", t) for t in near):
                print(f"  [고칠 것] {i}번째 표 위에 '표 n.' 제목이 없다 (앞에 온 글: {near[-2:]})")
                bad += 1
    n_fig = sum(1 for k, _ in seq if k == "그림파일")
    n_tab = sum(1 for k, _ in seq if k == "표")
    if not bad:
        print(f"  [맞음] 그림 {n_fig}개는 제목이 아래, 표 {n_tab}개는 제목이 위에 있다")


def check_hardcoded_refs() -> None:
    """본문이 그림·표를 **번호로 직접** 가리키고 있지 않은지 본다.

    번호는 등장 순서로 자동 부여되므로, 그림을 하나 끼워 넣으면 손으로 적은 번호가
    조용히 어긋난다. 실제로 인과 실험 그림을 넣자 '그림 5는 …'이 다른 그림을 가리키게 됐다.
    대신 <!--FIGREF:태그|조사--> 를 쓰면 번호와 조사가 함께 맞춰진다.

    표 제목("표 3. …")과 그림 제목("그림 2. …") 자체는 참조가 아니므로 뺀다.
    """
    print(chr(10) + "본문이 그림·표 번호를 손으로 적고 있지 않은가")
    hits = []
    pat = re.compile(r"(?<!<!--)(그림|표)\s(\d+)(?!\.)")
    for f in sorted(PAPER.glob("*.md")):
        if f.name.startswith("00_양식"):
            continue
        for ln, line in enumerate(f.read_text(encoding="utf-8").split(chr(10)), 1):
            t = line.strip()
            if t.startswith(("|", "<!--")):
                continue
            # 양식 규칙 자체를 설명하는 문장은 번호가 고정이다
            # ("학교 양식 지시에 따라 표 1은 임계 비용 λ*, 표 2는 실험 설정이다")
            if "양식" in t:
                continue
            for m in pat.finditer(t):
                hits.append((f.name, ln, m.group(0), t[max(0, m.start() - 18):m.end() + 18]))
    if not hits:
        print("  [맞음] 번호를 손으로 적은 곳이 없다 (전부 FIGREF/TABREF로 자동 부여)")
        return
    for name, ln, what, ctx in hits:
        print(f"  [고칠 것] {name}:{ln} '{what}' — …{ctx}…")
        print("        └ <!--FIGREF:태그|조사--> 로 바꿀 것 (번호와 조사가 자동으로 맞는다)")


def _plain(t: str) -> str:
    """굵게·코드·이스케이프 표시를 떼어 낸 비교용 문자열. 정규식을 쓰지 않는다 —
    문자 클래스 안의 역슬래시가 편집 과정에서 자주 깨졌다."""
    for ch in ("*", '`', '\\'):
        t = t.replace(ch, "")
    return t.strip()


def check_list_continuation() -> None:
    """목록 항목의 이어지는 줄이 별도 문단으로 떨어져 나오지 않았는지 확인한다.

    2026-08-29: 마크다운에서 여러 줄로 쓴 목록 항목의 둘째 줄부터가 별도 문단이 되어
    왼쪽 정렬 0으로 빠져 있었다. 화면상 '새 문단'처럼 보여 목록이 무너진다.

    확인 방법: 마크다운에서 **들여쓴 이어짐 줄**을 뽑아, 그 줄이 .docx에 **그 자체로**
    한 문단이 되어 있으면 쪼개진 것이다. 어림짐작이 아니라 원본과 결과를 직접 맞춘다.
    """
    print(chr(10) + "목록 항목이 쪼개지지 않았는가 (원본과 .docx 대조)")
    docx = ROOT / "졸업논문_초안v1.docx"
    if not docx.exists():
        print("  [건너뜀] .docx가 없다")
        return
    try:
        from docx import Document
    except Exception as e:
        print(f"  [확인 못 함] {e}")
        return

    cont = set()
    for f in sorted(PAPER.glob("*.md")):
        if f.name.startswith("00_양식"):
            continue
        lines = f.read_text(encoding="utf-8").split(chr(10))
        in_item = False
        for line in lines:
            t = line.strip()
            if not t:
                in_item = False
                continue
            is_item = t.startswith(("- ", "* ")) or re.match(r"^\d+\.\s", t)
            if is_item:
                in_item = True
                continue
            if in_item and line[:1].isspace():
                cont.add(_plain(t))
            else:
                in_item = False

    if not cont:
        print("  [맞음] 여러 줄로 쓴 목록 항목이 없다")
        return
    bad = []
    for para in Document(str(docx)).paragraphs:
        t = _plain(para.text)
        if t and t in cont:
            bad.append(t)
    if not bad:
        print(f"  [맞음] 이어짐 줄 {len(cont)}개가 모두 앞 항목에 합쳐져 있다")
        return
    for t in bad[:6]:
        print(f"  [고칠 것] 이어짐 줄이 따로 문단이 됐다: {t[:60]}…")


# 자동 생성되는 장은 환경마다 같은 문장을 쓰는 것이 정상이다(병렬 구조).
HAND_WRITTEN = ("00_초록.md", "01_서론.md", "02_관련연구.md", "03_방법.md", "05_결론.md")


def check_duplicate_sentences() -> None:
    """사람이 쓴 장에 같은 문장이 두 번 들어가 있지 않은지 본다.

    2026-08-29: Ⅲ장에 "두 가지 엄격도로 함께 보고한다."가 연달아 두 번 있었다.
    문단을 고쳐 쓰다 남은 흔적인데, 읽으면 바로 걸리는데도 눈으로는 잘 안 보인다.
    """
    print(chr(10) + "같은 문장이 두 번 들어가 있지 않은가 (사람이 쓴 장)")
    found = 0
    for name in HAND_WRITTEN:
        f = PAPER / name
        if not f.exists():
            continue
        sents = []
        for line in f.read_text(encoding="utf-8").split(chr(10)):
            t = line.strip()
            if not t or t.startswith(("|", "#", "<!--", ">", "-", "*")):
                continue
            for x in re.split(r"(?<=다[.])\s+", t):
                x = x.strip()
                if len(x) > 14:
                    sents.append(x)
        seen = {}
        for x in sents:
            seen[x] = seen.get(x, 0) + 1
        for x, n in seen.items():
            if n > 1:
                print(f"  [고칠 것] {name}: {n}회 반복 — {x[:58]}…")
                found += 1
    if not found:
        print("  [맞음] 사람이 쓴 장에 그대로 반복된 문장이 없다")


def main() -> None:
    print("=" * 78)
    print("원고 일관성 점검 — 사람이 판단할 목록 (자동 수정하지 않는다)")
    print("=" * 78)
    n_hits = 0
    # 참고문헌은 원고가 아니라 조립 스크립트에 들어 있다 — 거기 남은 표시도 함께 본다
    targets = sorted(PAPER.glob("*.md")) + [ROOT / "src" / "report" / "make_thesis_docx.py"]
    for f in targets:
        if f.name.startswith("00_양식") or not f.exists():
            continue
        text = f.read_text(encoding="utf-8")
        lines = text.splitlines()
        hits = []
        for name, pat, why in RISKS:
            for m in re.finditer(pat, text):
                ln = text[:m.start()].count("\n")
                hits.append((ln + 1, name, lines[ln].strip()[:100], why))
        if hits:
            print(f"\n[{f.name}]")
            for ln, name, ctx, why in sorted(hits):
                n_hits += 1
                print(f"  {ln:>4}행 ({name}) {ctx}")
                print(f"        └ {why}")
    print(f"\n확인할 곳 {n_hits}군데")
    check_dangling_headings()
    check_citation_order()
    check_english_case()
    check_source_control_chars()
    print(chr(10) + "손으로 옮겨 적은 표가 집계와 맞는가")
    try:
        from src.report.check_hand_numbers import check_switch_cost
        for line in check_switch_cost():
            print(line)
    except Exception as e:
        print(f"  [확인 못 함] {e}")
    check_ordinal_order()
    check_caption_position()
    check_hardcoded_refs()
    check_list_continuation()
    print(chr(10) + "논문이 적은 설정이 실제 config와 같은가")
    try:
        from src.report.check_paper_vs_config import main as _cfg_check
        import io
        import contextlib
        b = io.StringIO()
        with contextlib.redirect_stdout(b):
            _cfg_check()
        for line in b.getvalue().splitlines():
            t = line.strip()
            if t and not t.startswith("=") and "논문이 적은 설정이" not in t:
                print("  " + t)
    except Exception as e:
        print(f"  [확인 못 함] {e}")
    check_duplicate_sentences()
    print(f"\n절 참조가 맞는가")
    for line in section_refs():
        print(line)
    print(f"\n지시받은 그림·표 배치와 맞는가")
    for line in numbering_rules():
        print(line)
    print("\n현재 집계 기준 λ* (최강 규칙 포락선 대비) — 원고 수치와 대조할 것")
    for env, rows in numbers_in_use().items():
        bits = ", ".join(f"{k}: CI {v['CI']} / 점추정 {v['점추정']}" for k, v in rows.items())
        print(f"  {env}: {bits}")
    print("=" * 78)


if __name__ == "__main__":
    main()
