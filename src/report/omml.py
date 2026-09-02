"""워드 수식 객체(OMML)를 만든다.

왜 필요한가: 원고의 수식을 그냥 글자로 넣으면 워드에서 수식이 아니라 문자열이 된다.
심사에서 수식으로 보이지 않으면 논문답지 않고, 위첨자·아래첨자·합 기호가 제대로 서지 않는다.
그래서 실제 OMML(Office Math Markup Language) 객체로 넣는다.

이 파일은 이 논문에 나오는 **세 개의 수식만** 만든다. 범용 수식 변환기가 아니다 —
필요한 것만 손으로 정확히 짜는 편이 안전하다.

실행(자가 점검): python -m src.report.omml
"""
from __future__ import annotations

from docx.oxml import parse_xml
from docx.oxml.ns import qn

M = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _esc(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def r(t: str) -> str:
    """보통 글자."""
    return f"<m:r><m:t xml:space='preserve'>{_esc(t)}</m:t></m:r>"


def sup(base: str, e: str) -> str:
    """위첨자 — γ^i 처럼."""
    return f"<m:sSup><m:e>{base}</m:e><m:sup>{e}</m:sup></m:sSup>"


def sub(base: str, e: str) -> str:
    """아래첨자 — r_{t+i} 처럼."""
    return f"<m:sSub><m:e>{base}</m:e><m:sub>{e}</m:sub></m:sSub>"


def nary(chr_: str, lo: str, body: str) -> str:
    """합·곱 같은 큰 연산자. 위 첨자는 쓰지 않고 아래에만 조건을 단다."""
    return (
        "<m:nary><m:naryPr>"
        f"<m:chr m:val='{chr_}'/><m:limLoc m:val='undOvr'/>"
        "<m:supHide m:val='1'/>"
        "</m:naryPr>"
        f"<m:sub>{lo}</m:sub><m:sup/><m:e>{body}</m:e></m:nary>"
    )


def func(name: str, lo: str, body: str) -> str:
    """max 처럼 아래에 조건이 붙는 연산자."""
    return (
        "<m:func><m:funcPr><m:ctrlPr/></m:funcPr>"
        f"<m:fName><m:limLow><m:e>{r(name)}</m:e><m:lim>{lo}</m:lim></m:limLow></m:fName>"
        f"<m:e>{body}</m:e></m:func>"
    )


def frac(num: str, den: str) -> str:
    return f"<m:f><m:fPr><m:ctrlPr/></m:fPr><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def para(content: str, align: str = "center") -> str:
    """수식 하나를 담은 문단 XML. 가운데 정렬이 논문 관례다."""
    return (
        f"<w:p {W} {M}>"
        f"<w:pPr><w:jc w:val='{align}'/><w:spacing w:before='120' w:after='120' w:line='240' "
        "w:lineRule='auto'/></w:pPr>"
        f"<m:oMathPara><m:oMath>{content}</m:oMath></m:oMathPara>"
        "</w:p>"
    )


# ── 이 논문의 수식 세 개 ────────────────────────────────────────────────
def eq_qskip() -> str:
    """Q_skip(s,a,j) ← Σ_{i<j} γ^i r_{t+i} + γ^j max_{a'} Q(s_{t+j}, a')"""
    lhs = sub(r("Q"), r("skip")) + r("(s, a, j)") + r(" ← ")
    term1 = nary("∑", r("i < j"), sup(r("γ"), r("i")) + sub(r("r"), r("t+i")))
    inner = r("Q(") + sub(r("s"), r("t+j")) + r(", a′)")
    term2 = sup(r("γ"), r("j")) + func("max", r("a′"), inner)
    return lhs + term1 + r(" + ") + term2


def eq_cost() -> str:
    """r′(s, a) = r(s, a) − λ · 1[a ≠ a_noop]"""
    return (r("r′(s, a) = r(s, a) − λ · 1[a ≠ ") + sub(r("a"), r("noop")) + r("]"))


def eq_lambda_cross() -> str:
    """λ_교차 = (r_규칙 − r_무행동) / (에피소드당 과금 횟수)"""
    num = sub(r("r"), r("규칙")) + r(" − ") + sub(r("r"), r("무행동"))
    return sub(r("λ"), r("교차")) + r(" = ") + frac(num, r("에피소드당 과금 횟수"))


EQUATIONS = {
    "qskip": eq_qskip,
    "cost": eq_cost,
    "lambda_cross": eq_lambda_cross,
}


def insert(doc, key: str):
    """지금까지 쓴 본문 뒤에 수식 문단을 끼우고 그 문단 요소를 돌려준다.

    **body.append()를 쓰면 안 된다.** 문서 맨 끝에는 구역 설정 <w:sectPr>이 있고,
    규격상 그것이 본문의 마지막 요소여야 한다. append는 그 뒤에 붙어 버려
    워드·한워드가 파일 열기를 거부한다 (2026-09-02에 실제로 그래서 v2가 안 열렸다).
    다른 문단은 python-docx의 add_paragraph가 알아서 sectPr 앞에 넣는데,
    수식은 XML을 직접 만들어 넣으므로 여기서 같은 자리를 지켜야 한다.
    """
    if key not in EQUATIONS:
        raise KeyError(f"모르는 수식 이름: {key} (있는 것: {sorted(EQUATIONS)})")
    el = parse_xml(para(EQUATIONS[key]()))
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    if sect is None:
        body.append(el)
    else:
        sect.addprevious(el)
    return el


def main() -> int:
    """세 수식이 워드가 읽을 수 있는 XML로 만들어지는지 확인한다."""
    from docx import Document
    d = Document()
    for k in EQUATIONS:
        d.add_paragraph(k)
        insert(d, k)
    out = "results/reports/_omml_점검.docx"
    d.save(out)
    import zipfile
    x = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8")
    n = x.count("<m:oMath>")
    print(f"수식 {len(EQUATIONS)}개 작성 · 문서에 들어간 oMath {n}개 → "
          + ("정상" if n == len(EQUATIONS) else "확인 필요"))
    return 0 if n == len(EQUATIONS) else 1


if __name__ == "__main__":
    raise SystemExit(main())
