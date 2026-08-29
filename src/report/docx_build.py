"""학교 양식(붙임1_졸업논문양식.pdf)에 맞춘 .docx 조립 도구 — 저수준 부품.

양식 요구사항과 이 파일의 대응:
  · A4, 1단                      → section 설정
  · 더블 스페이싱                 → 문단 line_spacing = 2.0
  · 쪽 번호 하단 가운데 "- 1 -"   → 바닥글에 PAGE 필드
  · 장 번호는 로마자, 절은 아라비아 숫자 → 원고 마크다운에서 그대로 가져온다
  · 그림 제목은 그림 하단, 표 제목은 표 상단, 둘 다 국문·영문 병기 → add_figure / add_table
  · 인용 번호는 우측 상단 [n]     → 위 첨자 런으로 표시
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "바탕"          # 국문 본문 (윈도우 기본 명조)
HEAD_FONT = "맑은 고딕"      # 제목·캡션
BODY_PT = 10
LINE = 2.0                  # 더블 스페이싱


def set_run_font(run, name: str = BODY_FONT, size: float = BODY_PT,
                 bold: bool = False, italic: bool = False, color: str | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    return run


def new_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    # 기본 스타일도 맞춰 둔다 (빈 문단이 다른 글꼴로 튀지 않게)
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = st.paragraph_format
    pf.line_spacing = LINE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_page_number_footer(doc)
    return doc


def add_page_number_footer(doc: Document) -> None:
    """바닥글 가운데에 '- 1 -' 형식 쪽 번호. PAGE 필드를 쓰므로 Word가 자동으로 채운다."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run("- "), HEAD_FONT, 9)
    run = p.add_run()
    set_run_font(run, HEAD_FONT, 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(p.add_run(" -"), HEAD_FONT, 9)


def para(doc, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=BODY_PT,
         font=BODY_FONT, bold=False, line=LINE, before=0, after=0, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.widow_control = True        # 문단의 첫·마지막 한 줄이 혼자 다른 쪽에 남지 않게
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Pt(BODY_PT)   # 국문 논문 관례: 문단 첫 줄 한 글자 들여쓰기
    if text:
        add_rich_text(p, text, size=size, font=font, bold=bold)
    return p


CITE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")
TOKEN = re.compile(r"(\*\*[^*]+\*\*|\x60[^\x60]+\x60|\[\d+(?:\s*,\s*\d+)*\])")
BACKTICK = "\x60"


# 마크다운 뷰어용 역슬래시 이스케이프. 토큰을 나누기 전에 잠시 치워 둔다.
# (제어문자를 자리표시로 쓴다 — 원고에 나올 리 없는 글자여야 한다)
ESCAPES = {chr(92) + "*": chr(1), chr(92) + "_": chr(2), chr(92) + chr(96): chr(3)}


def add_rich_text(p, text: str, size=BODY_PT, font=BODY_FONT, bold=False):
    """굵게(**), 코드(백틱), [1] 인용(위 첨자)만 처리하는 최소 인라인 마크다운.

    원고에는 마크다운 뷰어용으로 λ\* 처럼 역슬래시로 escape한 별표가 있다. 그대로 두면
    **굵게** 안에 *가 섞여 토큰 나누기가 깨지므로, 먼저 치워 두었다가 마지막에 되돌린다.
    """
    for a, b in ESCAPES.items():
        text = text.replace(a, b)

    def unesc(x: str) -> str:
        for a, b in ESCAPES.items():
            x = x.replace(b, a[1:])
        return x

    for tok in TOKEN.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            set_run_font(p.add_run(unesc(tok[2:-2])), font, size, bold=True)
        elif tok.startswith(BACKTICK) and tok.endswith(BACKTICK) and len(tok) > 1:
            set_run_font(p.add_run(unesc(tok[1:-1])), "Consolas", size - 0.5)
        elif CITE.fullmatch(tok):
            r = p.add_run(tok)
            set_run_font(r, font, size, bold=bold)
            r.font.superscript = True      # 양식: 인용 위치 우측 상단에 [번호]
        else:
            set_run_font(p.add_run(unesc(tok)), font, size, bold=bold)
    return p


def keep_with_next(p, on: bool = True):
    """이 문단을 다음 문단과 붙여 둔다 — 쪽이 넘어가도 둘이 갈라지지 않는다.

    양식 요구: "그림과 표는 본문에서 언급된 쪽에 싣되, 다음 쪽으로 넘어가는 경우에는
    다음 쪽 맨 처음에 삽입한다." 제목만 앞 쪽에 남고 표가 뒤 쪽으로 넘어가면 이 규칙을 어긴다.
    """
    p.paragraph_format.keep_with_next = on
    return p


def no_widow(p):
    """문단의 첫 줄이나 마지막 줄이 혼자 다른 쪽에 남지 않게 한다."""
    p.paragraph_format.widow_control = True
    return p


def table_no_split(t, repeat_header: bool = True):
    """표가 쪽 경계에서 갈라지지 않게 한다.

    · cantSplit  : 한 줄(row)이 두 쪽에 걸쳐 잘리는 것을 막는다
    · tblHeader  : 표가 여러 쪽에 걸칠 만큼 길면 머리글 줄을 각 쪽에 다시 그린다
                   (머리글 없이 이어지는 표는 읽을 수 없다)
    · 마지막 줄을 뺀 모든 줄에 keep_with_next — 짧은 표는 통째로 다음 쪽으로 밀린다
    """
    rows = t.rows
    for i, r in enumerate(rows):
        trPr = r._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
        if i == 0 and repeat_header:
            trPr.append(OxmlElement("w:tblHeader"))
        if i < len(rows) - 1:
            for c in r.cells:
                for par in c.paragraphs:
                    par.paragraph_format.keep_with_next = True
    return t


def add_chapter_heading(doc, text: str):
    return keep_with_next(para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, font=HEAD_FONT,
                               bold=True, line=1.3, before=16, after=6, first_indent=False))


def add_section_heading(doc, text: str):
    return keep_with_next(para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, font=HEAD_FONT,
                               bold=True, line=1.3, before=10, after=4, first_indent=False))


def add_sub_heading(doc, text: str):
    return keep_with_next(para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10, font=HEAD_FONT,
                               bold=True, line=1.3, before=8, after=3, first_indent=False))


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE
    pf.left_indent = Cm(0.6 + 0.5 * level)
    pf.first_line_indent = Cm(-0.35)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_rich_text(p, ("· " if level == 0 else "- ") + text)
    return p


def add_numbered(doc, text: str):
    """이미 '1.' 같은 번호가 붙은 줄 — 글머리표를 덧붙이지 않고 매달린 들여쓰기만 준다."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE
    pf.left_indent = Cm(0.6)
    pf.first_line_indent = Cm(-0.6)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_rich_text(p, text)
    return p


def add_figure(doc, image_path: Path, ko: str, en: str, note: str = "",
               width_cm: float = 15.0):
    """그림 + 하단 캡션(국문·영문 병기). 양식: 그림 제목은 그림 하단."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    keep_with_next(p)                      # 그림과 그 아래 제목이 갈라지지 않게
    cap = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, before=4, after=2,
               first_indent=False)
    set_run_font(cap.add_run(ko), HEAD_FONT, 9.5, bold=True)
    keep_with_next(cap)
    cap2 = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, before=0, after=4,
                first_indent=False)
    set_run_font(cap2.add_run(en), HEAD_FONT, 9, italic=True)
    if note:
        n = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, before=0, after=8,
                 first_indent=False)
        set_run_font(n.add_run(note), HEAD_FONT, 8.5, color="555555")
    return p


def add_table(doc, header: list[str], rows: list[list[str]], ko: str, en: str,
              note: str = "", widths: list[float] | None = None):
    """표 제목(국문·영문)을 위에 붙이고 표를 그린다. 양식: 표 제목은 표 상단."""
    cap = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, before=10, after=2,
               first_indent=False)
    set_run_font(cap.add_run(ko), HEAD_FONT, 9.5, bold=True)
    keep_with_next(cap)                    # 양식: 표 제목은 표 상단 — 갈라지면 안 된다
    cap2 = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, before=0, after=3,
                first_indent=False)
    set_run_font(cap2.add_run(en), HEAD_FONT, 9, italic=True)
    keep_with_next(cap2)

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        # 머리글 칸도 인라인 마크다운을 거친다 — 원고 표 머리글에 **굵게**가 들어가는 경우가 있다
        add_rich_text(p, str(h), size=9, font=HEAD_FONT, bold=True)
        _shade(c, "EDEDED")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            if i >= len(cells):
                continue
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            add_rich_text(p, str(v), size=9, font=BODY_FONT)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                if i < len(r.cells):
                    r.cells[i].width = Cm(w)
    table_no_split(t)
    if note:
        n = para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, before=3, after=8,
                 first_indent=False)
        set_run_font(n.add_run(note), HEAD_FONT, 8.5, color="555555")
    else:
        para(doc, "", line=1.0, after=6, first_indent=False)
    return t


def _shade(cell, hexcolor: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p
